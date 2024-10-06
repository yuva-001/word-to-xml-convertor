from flask import Flask, request, render_template, redirect, flash, jsonify
import docx
from lxml import etree
import base64
import os
from PIL import Image
from io import BytesIO
import re
from werkzeug.utils import secure_filename
import platform
import subprocess

app = Flask(__name__)
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['CONVERTED_FOLDER'] = 'converted'
app.config['SECRET_KEY'] = 'supersecretkey'

# Ensure upload and converted directories exist
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
os.makedirs(app.config['CONVERTED_FOLDER'], exist_ok=True)


@app.route('/', methods=['GET', 'POST'])
def index():
    if request.method == 'POST':
        if 'files' not in request.files:
            flash('No file part')
            return redirect(request.url)
        
        files = request.files.getlist('files')
        converted_files = []

        for file in files:
            if file.filename == '':
                flash('No selected file')
                return redirect(request.url)
            
            if file:
                filename = secure_filename(file.filename)
                docx_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
                file.save(docx_path)
                xml_filename = f"{os.path.splitext(filename)[0]}.xml"
                xml_path = os.path.join(app.config['CONVERTED_FOLDER'], xml_filename)
                convert_docx_to_xml(docx_path, xml_path)
                converted_files.append(xml_filename)
        
        return render_template('index.html', conversions=converted_files)

    return render_template('index.html', conversions=[])


@app.route('/upload', methods=['POST'])
def upload():
    if 'files' not in request.files:
        flash('No file part')
        return redirect(request.url)
    
    files = request.files.getlist('files')
    converted_files = []

    for file in files:
        if file.filename == '':
            flash('No selected file')
            return redirect(request.url)
        
        if file:
            filename = secure_filename(file.filename)
            docx_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(docx_path)
            xml_filename = f"{os.path.splitext(filename)[0]}.xml"
            xml_path = os.path.join(app.config['CONVERTED_FOLDER'], xml_filename)
            convert_docx_to_xml(docx_path, xml_path)
            converted_files.append(xml_filename)
    
    return render_template('index.html', conversions=converted_files)

@app.route('/open-folder')
def open_folder():
    folder_path = os.path.abspath(app.config['CONVERTED_FOLDER'])
    if platform.system() == "Windows":
        os.startfile(folder_path)
    elif platform.system() == "Darwin":
        subprocess.Popen(["open", folder_path])
    else:
        subprocess.Popen(["xdg-open", folder_path])
    return '', 204

@app.route('/delete/<filename>', methods=['POST'])
def delete_file(filename):
    file_path = os.path.join(app.config['CONVERTED_FOLDER'], filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        return jsonify({'status': 'success'})
    else:
        return jsonify({'status': 'error'})


def convert_docx_to_xml(docx_path, xml_path):
    doc = docx.Document(docx_path)
    root = etree.Element("document")

    main_section = etree.SubElement(root, "section", id="s0")
    section_stack = [(main_section, 0)]
    section_count = 0
    current_list = None
    first_paragraph_processed = False

    def create_section(parent, level):
        nonlocal section_count
        section_count += 1
        section_elem = etree.Element("section", id=f"s{section_count}", level=str(level))
        parent.append(section_elem)
        return section_elem

    def add_paragraph(paragraph, parent, is_first_paragraph):
        if paragraph.text.strip():
            if is_first_paragraph:
                para_elem = etree.SubElement(parent, "art_title")
            else:
                para_elem = etree.SubElement(parent, "paragraph")

            current_elem = para_elem
            for run in paragraph.runs:
                text = run.text
                if text:
                    if run.bold:
                        if current_elem.tag == 'b':
                            current_elem.text = (current_elem.text or "") + text
                        else:
                            current_elem = etree.SubElement(para_elem, "b")
                            current_elem.text = text
                    elif run.italic:
                        if current_elem.tag == 'i':
                            current_elem.text = (current_elem.text or "") + text
                        else:
                            current_elem = etree.SubElement(para_elem, "i")
                            current_elem.text = text
                    else:
                        if current_elem.tag == 'paragraph':
                            current_elem.text = (current_elem.text or "") + text
                        else:
                            current_elem = para_elem
                            current_elem.text = (current_elem.text or "") + text

    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            heading_level = int(re.search(r'\d+', paragraph.style.name).group())
            while section_stack and section_stack[-1][1] >= heading_level:
                section_stack.pop()
            current_section, _ = section_stack[-1]
            new_section = create_section(current_section, heading_level)
            section_stack.append((new_section, heading_level))
            add_paragraph(paragraph, new_section, first_paragraph_processed)
        else:
            if first_paragraph_processed:
                add_paragraph(paragraph, section_stack[-1][0], False)
            else:
                add_paragraph(paragraph, main_section, True)
                first_paragraph_processed = True

    def encode_image(img):
        img_stream = BytesIO()
        img.save(img_stream, format='PNG')
        img_str = base64.b64encode(img_stream.getvalue()).decode('utf-8')
        return img_str

    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            try:
                img = Image.open(doc.part.rels[rel.target_ref].target_part.blob)
                img_elem = etree.SubElement(section_stack[-1][0], "image")
                img_elem.text = encode_image(img)
            except KeyError:
                print(f"KeyError: {rel.target_ref} not found.")

    tree = etree.ElementTree(root)
    tree.write(xml_path, pretty_print=True, xml_declaration=True, encoding="UTF-8")


if __name__ == "__main__":
    app.run(debug=True)
